from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement

from convert_to_md.mathutil import qn_m
from convert_to_md.pipeline import convert_file


def r_text(s: str):
    r = OxmlElement("m:r")
    t = OxmlElement("m:t")
    t.text = s
    r.append(t)
    return r


def omath_para(builder):
    o_math_para = OxmlElement("m:oMathPara")
    o_math = OxmlElement("m:oMath")
    builder(o_math)
    o_math_para.append(o_math)
    return o_math_para


def main() -> None:
    raw = Path("samples/raw")
    out_dir = Path("samples/out")
    doc = Document()
    doc.add_heading("DOCX Formula Sample", level=1)

    doc.add_paragraph("Energy equation:")
    p = doc.add_paragraph()

    def emc(o):
        for ch in ["E", "=", "m"]:
            o.append(r_text(ch))
        s_sup = OxmlElement("m:sSup")
        e = OxmlElement("m:e")
        e.append(r_text("c"))
        s_sup.append(e)
        sup = OxmlElement("m:sup")
        sup.append(r_text("2"))
        s_sup.append(sup)
        o.append(s_sup)

    p._p.append(omath_para(emc))

    doc.add_paragraph("Definite integral:")
    p = doc.add_paragraph()

    def integ(o):
        nary = OxmlElement("m:nary")
        nary_pr = OxmlElement("m:naryPr")
        chr_ = OxmlElement("m:chr")
        chr_.set(qn_m("val"), "∫")
        nary_pr.append(chr_)
        nary.append(nary_pr)
        sub = OxmlElement("m:sub")
        sub.append(r_text("0"))
        nary.append(sub)
        sup = OxmlElement("m:sup")
        sup.append(r_text("1"))
        nary.append(sup)
        e = OxmlElement("m:e")
        e.append(r_text("x"))
        nary.append(e)
        o.append(nary)
        o.append(r_text(" dx"))

    p._p.append(omath_para(integ))

    doc.add_paragraph("Matrix:")
    p = doc.add_paragraph()

    def mat(o):
        m = OxmlElement("m:m")
        for row in (("1", "2"), ("3", "4")):
            mr = OxmlElement("m:mr")
            for val in row:
                ee = OxmlElement("m:e")
                ee.append(r_text(val))
                mr.append(ee)
            m.append(mr)
        o.append(m)

    p._p.append(omath_para(mat))
    doc.add_paragraph("Inline text after equations.")

    path = raw / "formulas.docx"
    doc.save(str(path))
    print("saved", path)
    out = convert_file(path, out_dir, use_cache=False)
    print(out.read_text(encoding="utf-8"))


if __name__ == "__main__":
    main()
