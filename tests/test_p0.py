from __future__ import annotations

import json
from pathlib import Path

import pytest

from convert_to_md.detect import sniff
from convert_to_md.pipeline import convert_file
from convert_to_md.render.markdown import render_markdown
from convert_to_md.ir import DocumentIR


FIXTURES = Path(__file__).parent / "fixtures"


def test_sniff_json(tmp_path: Path):
    p = tmp_path / "a.json"
    p.write_text('{"x": 1}', encoding="utf-8")
    assert sniff(p).kind == "json"


def test_convert_json(tmp_path: Path):
    src = tmp_path / "sample.json"
    src.write_text(json.dumps({"hello": "world", "n": 1}, ensure_ascii=False), encoding="utf-8")
    out = convert_file(src, tmp_path)
    assert out.name == "sample.json.md"
    text = out.read_text(encoding="utf-8")
    assert "```json" in text
    assert "hello" in text


def test_convert_csv(tmp_path: Path):
    src = tmp_path / "t.csv"
    src.write_text("name,age\nAlice,30\nBob,25\n", encoding="utf-8")
    out = convert_file(src, tmp_path)
    text = out.read_text(encoding="utf-8")
    assert "| name | age |" in text
    assert "Alice" in text


def test_convert_html(tmp_path: Path):
    src = tmp_path / "p.html"
    src.write_text(
        "<html><head><title>T</title></head><body><h1>Hi</h1><p>para</p></body></html>",
        encoding="utf-8",
    )
    out = convert_file(src, tmp_path)
    text = out.read_text(encoding="utf-8")
    assert "Hi" in text
    assert "para" in text


def test_html_strips_nav_noise(tmp_path: Path):
    src = tmp_path / "wikiish.html"
    src.write_text(
        """
        <html><head><title>Markdown - Wikipedia</title></head>
        <body>
          <nav class="vector-main-menu" role="navigation">
            <a href="/wiki/Main_Page">Main page</a>
            <a href="/wiki/Contents">Contents</a>
            <a href="/wiki/Random">Random article</a>
          </nav>
          <header class="vector-header">Top chrome</header>
          <main>
            <div id="mw-content-text">
              <div class="mw-parser-output">
                <h1>Markdown</h1>
                <p>Markdown is a lightweight markup language.<sup class="reference"><a href="#cite">[1]</a></sup></p>
                <table class="infobox"><tr><th>Type</th><td>Markup</td></tr></table>
                <div class="navbox"><a href="/a">A</a><a href="/b">B</a><a href="/c">C</a><a href="/d">D</a><a href="/e">E</a></div>
              </div>
            </div>
          </main>
          <footer>Site footer</footer>
        </body></html>
        """,
        encoding="utf-8",
    )
    out = convert_file(src, tmp_path, use_cache=False)
    text = out.read_text(encoding="utf-8")
    assert "lightweight markup language" in text
    assert "Random article" not in text
    assert "Site footer" not in text
    assert "Top chrome" not in text
    assert "infobox" not in text.lower()
    assert "| Type |" not in text
    assert "[1]" not in text


def test_convert_txt(tmp_path: Path):
    src = tmp_path / "n.txt"
    src.write_text("line1\n\nline2\n", encoding="utf-8")
    out = convert_file(src, tmp_path)
    text = out.read_text(encoding="utf-8")
    assert "line1" in text
    assert "line2" in text


def test_convert_xml(tmp_path: Path):
    src = tmp_path / "d.xml"
    src.write_text("<root><a>1</a></root>", encoding="utf-8")
    out = convert_file(src, tmp_path)
    text = out.read_text(encoding="utf-8")
    assert "```xml" in text
    assert "<root>" in text


def test_convert_docx(tmp_path: Path):
    pytest.importorskip("docx")
    from docx import Document

    src = tmp_path / "doc.docx"
    d = Document()
    d.add_heading("Title One", level=1)
    d.add_paragraph("Hello docx")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    d.save(str(src))

    out = convert_file(src, tmp_path)
    text = out.read_text(encoding="utf-8")
    assert "Title One" in text
    assert "Hello docx" in text
    assert "A" in text


def test_convert_xlsx(tmp_path: Path):
    pytest.importorskip("openpyxl")
    from openpyxl import Workbook

    src = tmp_path / "s.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "SheetA"
    ws.append(["x", "y"])
    ws.append([1, 2])
    wb.save(str(src))

    out = convert_file(src, tmp_path)
    text = out.read_text(encoding="utf-8")
    assert "SheetA" in text
    assert "x" in text


def test_cache_skips_second_run(tmp_path: Path):
    src = tmp_path / "c.json"
    src.write_text('{"a":1}', encoding="utf-8")
    out1 = convert_file(src, tmp_path, use_cache=True)
    mtime1 = out1.stat().st_mtime_ns
    out2 = convert_file(src, tmp_path, use_cache=True)
    assert out1 == out2
    assert out2.stat().st_mtime_ns == mtime1


def test_pdf_authors_not_merged(tmp_path: Path):
    fitz = pytest.importorskip("fitz")
    src = tmp_path / "authors.pdf"
    doc = fitz.open()
    page = doc.new_page(width=612, height=792)
    page.insert_text((200, 120), "Attention Is All You Need", fontsize=18)
    # side-by-side author cards
    page.insert_text((80, 200), "Ada Lovelace*", fontsize=9)
    page.insert_text((80, 214), "Analytical Engine Lab", fontsize=9)
    page.insert_text((80, 228), "ada@example.com", fontsize=9)
    page.insert_text((250, 200), "Alan Turing*", fontsize=9)
    page.insert_text((250, 214), "Bletchley Research", fontsize=9)
    page.insert_text((250, 228), "alan@example.com", fontsize=9)
    page.insert_text((420, 200), "Grace Hopper*", fontsize=9)
    page.insert_text((420, 214), "Navy Research", fontsize=9)
    page.insert_text((420, 228), "grace@example.com", fontsize=9)
    page.insert_text((250, 300), "Abstract", fontsize=12)
    page.insert_text((100, 330), "This paper studies sequence transduction models.", fontsize=10)
    doc.save(str(src))
    doc.close()

    out = convert_file(src, tmp_path, use_cache=False)
    text = out.read_text(encoding="utf-8")
    assert "Authors" in text
    assert "Ada Lovelace" in text
    assert "Alan Turing" in text
    # authors should not be smashed into one long run without separators
    assert "Ada Lovelace* Alan Turing*" not in text
    assert "Abstract" in text


def test_pdf_no_page_headings_by_default(tmp_path: Path):
    fitz = pytest.importorskip("fitz")
    src = tmp_path / "multi.pdf"
    doc = fitz.open()
    for i, title in enumerate(["Intro", "Method"], 1):
        page = doc.new_page()
        page.insert_text((72, 72), title, fontsize=16)
        page.insert_text((72, 110), f"Body text for page {i}.", fontsize=11)
    doc.save(str(src))
    doc.close()

    out = convert_file(src, tmp_path, use_cache=False, pdf_page_headings=False)
    text = out.read_text(encoding="utf-8")
    assert "## Page 1" not in text
    assert "Intro" in text
    assert "Method" in text

    out2 = convert_file(
        src,
        tmp_path / "with_pages.md",
        use_cache=False,
        pdf_page_headings=True,
    )
    text2 = out2.read_text(encoding="utf-8")
    assert "## Page 1" in text2


def test_pdf_ocr_off_on_empty_scan_message(tmp_path: Path):
    fitz = pytest.importorskip("fitz")
    src = tmp_path / "blankish.pdf"
    doc = fitz.open()
    doc.new_page()  # empty page, no text layer
    doc.save(str(src))
    doc.close()

    out = convert_file(src, tmp_path, use_cache=False, pdf_ocr="off")
    text = out.read_text(encoding="utf-8")
    assert "scanned PDF" in text.lower() or "No extractable text" in text


def test_render_empty_assets_not_created(tmp_path: Path):
    doc = DocumentIR(title="t")
    doc.paragraph("only text")
    out = tmp_path / "o.md"
    assets = tmp_path / "o_assets"
    from convert_to_md.render.markdown import write_markdown

    write_markdown(doc, out, assets_dir=assets)
    assert out.exists()
    assert not assets.exists()


def test_sniff_image(tmp_path: Path):
    # minimal PNG header
    p = tmp_path / "x.png"
    p.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 20)
    assert sniff(p).kind == "image"


def test_parallel_convert_path(tmp_path: Path):
    from convert_to_md.pipeline import convert_path

    src_dir = tmp_path / "in"
    out_dir = tmp_path / "out"
    src_dir.mkdir()
    for i in range(4):
        (src_dir / f"a{i}.json").write_text(f'{{"i": {i}}}', encoding="utf-8")
    paths = convert_path(src_dir, out_dir, workers=2, use_cache=False)
    assert len(paths) == 4
    assert all(p.suffix == ".md" for p in paths)


def test_epub_boilerplate_helpers():
    from convert_to_md.converters.epub import _cleanup_md, _is_boilerplate_chapter

    dirty = "*** START OF THE PROJECT GUTENBERG EBOOK X ***\n\nHello chapter body.\n"
    assert "START OF THE PROJECT GUTENBERG" not in _cleanup_md(dirty)
    assert "Hello chapter body" in _cleanup_md(dirty)
    license_page = (
        "This eBook is for the use of anyone anywhere in the United States and "
        "most other parts of the world at no cost and with almost no restrictions "
        "whatsoever. " * 3
    )
    assert _is_boilerplate_chapter(license_page, "Project Gutenberg")
    story = '“My dear Mr. Bennet,” said his lady to him one day, “have you heard?”\n' * 5
    assert not _is_boilerplate_chapter(story, "Chapter I")


def test_image_and_pdf_ocr_if_tesseract(tmp_path: Path):
    from convert_to_md.ocr import available_engines

    if "tesseract" not in available_engines():
        pytest.skip("tesseract not available")

    from PIL import Image, ImageDraw
    import fitz

    # image OCR
    img_path = tmp_path / "hello.png"
    im = Image.new("RGB", (400, 120), "white")
    d = ImageDraw.Draw(im)
    d.text((20, 40), "Hello OCR World", fill="black")
    im.save(img_path)
    out = convert_file(img_path, tmp_path, use_cache=False)
    text = out.read_text(encoding="utf-8")
    assert "Hello" in text or "OCR" in text or "World" in text

    # scanned-like PDF via rendered text page without text layer:
    # create PDF with only image content
    pdf_path = tmp_path / "scanlike.pdf"
    doc = fitz.open()
    page = doc.new_page(width=400, height=120)
    page.insert_image(page.rect, filename=str(img_path))
    doc.save(str(pdf_path))
    doc.close()
    out_pdf = convert_file(pdf_path, tmp_path, use_cache=False, pdf_ocr="force")
    pdf_text = out_pdf.read_text(encoding="utf-8")
    assert "Hello" in pdf_text or "OCR" in pdf_text or "World" in pdf_text


def test_html_mathml_and_latex(tmp_path: Path):
    src = tmp_path / "m.html"
    src.write_text(
        """
        <html><head><title>Math</title></head><body>
        <h1>Math</h1>
        <p>Inline
        <math xmlns="http://www.w3.org/1998/Math/MathML">
          <semantics>
            <mrow><mi>E</mi><mo>=</mo><mi>m</mi><msup><mi>c</mi><mn>2</mn></msup></mrow>
            <annotation encoding="application/x-tex">E=mc^2</annotation>
          </semantics>
        </math>
        and display:</p>
        <math xmlns="http://www.w3.org/1998/Math/MathML" display="block">
          <semantics>
            <mrow><mi>a</mi><mo>+</mo><mi>b</mi></mrow>
            <annotation encoding="application/x-tex">a+b</annotation>
          </semantics>
        </math>
        <script type="math/tex; mode=display">\\sum_{i=1}^n i</script>
        </body></html>
        """,
        encoding="utf-8",
    )
    out = convert_file(src, tmp_path, use_cache=False)
    text = out.read_text(encoding="utf-8")
    assert "$E=mc^2$" in text or "$E = mc^2$" in text or "E=mc^2" in text or "E=mc^{2}" in text
    assert "$$" in text
    assert "a+b" in text
    assert "\\sum" in text or "sum" in text


def test_docx_omml_formula(tmp_path: Path):
    from docx import Document
    from docx.oxml import OxmlElement

    src = tmp_path / "eq.docx"
    d = Document()
    d.add_heading("Eq", level=1)
    p = d.add_paragraph()
    oMathPara = OxmlElement("m:oMathPara")
    oMath = OxmlElement("m:oMath")
    for ch in ["E", "=", "m"]:
        r = OxmlElement("m:r")
        t = OxmlElement("m:t")
        t.text = ch
        r.append(t)
        oMath.append(r)
    sSup = OxmlElement("m:sSup")
    e = OxmlElement("m:e")
    r = OxmlElement("m:r")
    t = OxmlElement("m:t")
    t.text = "c"
    r.append(t)
    e.append(r)
    sSup.append(e)
    sup = OxmlElement("m:sup")
    r2 = OxmlElement("m:r")
    t2 = OxmlElement("m:t")
    t2.text = "2"
    r2.append(t2)
    sup.append(r2)
    sSup.append(sup)
    oMath.append(sSup)
    oMathPara.append(oMath)
    p._p.append(oMathPara)
    d.save(str(src))

    out = convert_file(src, tmp_path, use_cache=False)
    text = out.read_text(encoding="utf-8")
    assert "E" in text and "c" in text
    assert "$$" in text or "$" in text


def test_pdf_unicode_formula(tmp_path: Path):
    fitz = pytest.importorskip("fitz")
    src = tmp_path / "f.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "E = mc²", fontsize=14)
    page.insert_text((72, 110), "Normal sentence without formula symbols here.", fontsize=11)
    doc.save(str(src))
    doc.close()
    out = convert_file(src, tmp_path, use_cache=False)
    text = out.read_text(encoding="utf-8")
    assert "$$" in text or "mc" in text


def test_html_strips_word_joiners_and_bullet_dots(tmp_path: Path):
    src = tmp_path / "join.html"
    src.write_text(
        """
        <html><head><title>[1909.03550] Optimization Notes All rights reserved.</title></head>
        <body>
          <h1 class="ltx_title_document">Optimization Notes All rights reserved.</h1>
          <ul><li>• Euclidean space</li></ul>
          <p>y = ⁠1/2⁠ x</p>
        </body></html>
        """,
        encoding="utf-8",
    )
    out = convert_file(src, tmp_path, use_cache=False)
    text = out.read_text(encoding="utf-8")
    assert "All rights reserved" not in text
    assert "[1909" not in text
    assert "⁠" not in text
    assert "- •" not in text
    assert "Euclidean space" in text


def test_pretty_aligned_latex():
    from convert_to_md.mathutil import latex_to_markdown, pretty_latex

    raw = r"{\begin{aligned}a&=1\\[3mu]b&=2\end{aligned}}"
    pretty = pretty_latex(raw)
    assert "\\begin{aligned}" in pretty
    assert "\\\\\n" in pretty
    assert "vphantom" not in pretty
    md = latex_to_markdown(raw, display=True)
    assert md.startswith("$$")
    assert "a&=1" in md and "b&=2" in md

    # nested $ must not break display fences
    nested = r"\mbox{person $i$ likes}"
    md2 = latex_to_markdown(nested, display=True)
    assert md2.count("$$") == 2
    assert "\\$i\\$" in md2 or r"\$i\$" in md2


def test_postprocess_formula_latex():
    from convert_to_md.mathutil import postprocess_formula_latex

    assert postprocess_formula_latex("E=m c^{2}") == "E=mc^{2}"
    assert "underline" not in postprocess_formula_latex(r"x=\underline{{ac}}")


def test_omml_integral_and_matrix():
    from docx.oxml import OxmlElement
    from convert_to_md.mathutil import omml_element_to_latex, qn_m

    # integral: ∫_0^1 x dx
    o = OxmlElement("m:oMath")
    nary = OxmlElement("m:nary")
    naryPr = OxmlElement("m:naryPr")
    chr_ = OxmlElement("m:chr")
    chr_.set(qn_m("val"), "∫")
    naryPr.append(chr_)
    nary.append(naryPr)
    sub = OxmlElement("m:sub")
    r = OxmlElement("m:r"); t = OxmlElement("m:t"); t.text = "0"; r.append(t); sub.append(r)
    nary.append(sub)
    sup = OxmlElement("m:sup")
    r = OxmlElement("m:r"); t = OxmlElement("m:t"); t.text = "1"; r.append(t); sup.append(r)
    nary.append(sup)
    e = OxmlElement("m:e")
    r = OxmlElement("m:r"); t = OxmlElement("m:t"); t.text = "x"; r.append(t); e.append(r)
    nary.append(e)
    o.append(nary)
    latex = omml_element_to_latex(o)
    assert "\\int" in latex
    assert "0" in latex and "1" in latex

    # matrix 1 2 / 3 4
    o2 = OxmlElement("m:oMath")
    m = OxmlElement("m:m")
    for a, b in (("1", "2"), ("3", "4")):
        mr = OxmlElement("m:mr")
        for val in (a, b):
            ee = OxmlElement("m:e")
            rr = OxmlElement("m:r"); tt = OxmlElement("m:t"); tt.text = val; rr.append(tt); ee.append(rr)
            mr.append(ee)
        m.append(mr)
    o2.append(m)
    latex2 = omml_element_to_latex(o2)
    assert "\\begin{matrix}" in latex2
    assert "1 & 2" in latex2
    assert "3 & 4" in latex2


def test_pdf_formula_ocr_on_equation_image(tmp_path: Path):
    from convert_to_md.ocr import available_engines

    if "tesseract" not in available_engines():
        pytest.skip("tesseract not available")

    from PIL import Image, ImageDraw
    import fitz

    # wide short formula-like image (no text layer)
    img_path = tmp_path / "eq.png"
    im = Image.new("RGB", (520, 80), "white")
    d = ImageDraw.Draw(im)
    d.text((20, 28), "E = mc^2", fill="black")
    im.save(img_path)

    pdf_path = tmp_path / "eq_image.pdf"
    doc = fitz.open()
    page = doc.new_page(width=612, height=200)
    # place as a short-wide image strip
    page.insert_image(fitz.Rect(80, 60, 520, 130), filename=str(img_path))
    doc.save(str(pdf_path))
    doc.close()

    out = convert_file(
        pdf_path,
        tmp_path,
        use_cache=False,
        pdf_ocr="off",
        pdf_formula_ocr="force",
        formula_ocr_engine="tesseract",
    )
    text = out.read_text(encoding="utf-8")
    # keep image asset and OCR formula attempt
    assert "image" in text.lower() or "![" in text or "$$" in text or "E" in text
    # with force + tesseract, OCR text should surface as formula-ish content
    assert "E" in text and ("mc" in text.lower() or "m" in text)


def test_golden_html_math_snippet(tmp_path: Path):
    src = tmp_path / "g.html"
    src.write_text(
        """
        <html><head><title>Golden Math</title></head><body>
        <h1>Golden Math</h1>
        <p>Energy
        <math xmlns="http://www.w3.org/1998/Math/MathML" display="block">
          <semantics>
            <mrow><mi>E</mi><mo>=</mo><mi>m</mi><msup><mi>c</mi><mn>2</mn></msup></mrow>
            <annotation encoding="application/x-tex">E=mc^{2}</annotation>
          </semantics>
        </math>
        </p>
        </body></html>
        """,
        encoding="utf-8",
    )
    out = convert_file(src, tmp_path, use_cache=False)
    text = out.read_text(encoding="utf-8")
    assert "Golden Math" in text
    assert "$$" in text
    assert "E=mc^{2}" in text or "E=mc" in text


def test_config_env_override(monkeypatch, tmp_path: Path):
    from convert_to_md.config import load_config

    monkeypatch.chdir(tmp_path)
    monkeypatch.setenv("CONVERT_TO_MD_WORKERS", "3")
    monkeypatch.setenv("CONVERT_TO_MD_PDF_OCR", "off")
    monkeypatch.setenv("CONVERT_TO_MD_FORMULA_OCR_ENGINE", "tesseract")
    cfg = load_config(tmp_path)
    assert cfg.workers == 3
    assert cfg.pdf_ocr == "off"
    assert cfg.formula_ocr_engine == "tesseract"


def test_config_toml_file(tmp_path: Path, monkeypatch):
    from convert_to_md.config import load_config

    monkeypatch.chdir(tmp_path)
    (tmp_path / "convert-to-md.toml").write_text(
        """
[convert-to-md]
workers = 2
pdf_formula_ocr = "force"
formula_ocr_engine = "pix2tex"
use_cache = false
""".strip(),
        encoding="utf-8",
    )
    cfg = load_config(tmp_path)
    assert cfg.workers == 2
    assert cfg.pdf_formula_ocr == "force"
    assert cfg.formula_ocr_engine == "pix2tex"
    assert cfg.use_cache is False


def test_formula_ocr_fallback_tesseract(tmp_path: Path):
    from convert_to_md.formula_ocr import available_formula_engines, ocr_formula_path
    from convert_to_md.ocr import available_engines

    if "tesseract" not in available_engines():
        pytest.skip("tesseract not available")

    from PIL import Image, ImageDraw

    img = tmp_path / "eq.png"
    im = Image.new("RGB", (420, 90), "white")
    d = ImageDraw.Draw(im)
    d.text((20, 30), "E = mc^2", fill="black")
    im.save(img)

    engines = available_formula_engines()
    assert "tesseract" in engines
    result = ocr_formula_path(img, engine="tesseract")
    assert result.engine == "tesseract"
    assert result.latex
    assert "E" in result.latex


def test_html_equation_number_attached(tmp_path: Path):
    src = tmp_path / "eqnum.html"
    src.write_text(
        """
        <html><body>
        <table>
          <tr class="ltx_equation">
            <td>
              <math xmlns="http://www.w3.org/1998/Math/MathML" display="block">
                <semantics>
                  <mrow><mi>E</mi><mo>=</mo><mi>m</mi><msup><mi>c</mi><mn>2</mn></msup></mrow>
                  <annotation encoding="application/x-tex">E=mc^{2}</annotation>
                </semantics>
              </math>
            </td>
            <td><span class="ltx_tag ltx_tag_equation">(1.1)</span></td>
          </tr>
        </table>
        </body></html>
        """,
        encoding="utf-8",
    )
    out = convert_file(src, tmp_path, use_cache=False)
    text = out.read_text(encoding="utf-8")
    assert "E=mc^{2}" in text or "E=mc" in text
    assert "*(1.1)*" in text
    # original bare number should not remain as a lone table cell artifact ideally
    assert text.count("(1.1)") == 1


def test_attach_orphan_equation_numbers():
    from convert_to_md.mathutil import attach_orphan_equation_numbers, cleanup_formula_table_chrome

    raw = "$$\na+b=c\n$$\n(2.3)\n"
    out = attach_orphan_equation_numbers(raw)
    assert "*(2.3)*" in out

    messy = "   |  |  |  |\n   | --- | --- | --- |\n   |  | $$\nx=1\n$$ |  |\n"
    cleaned = cleanup_formula_table_chrome(messy)
    assert "$$" in cleaned
    assert "x=1" in cleaned
    assert "---" not in cleaned
    assert cleaned.count("|") < messy.count("|")



